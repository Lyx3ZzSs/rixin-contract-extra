"""PaddleOCR provider — text direct extraction + PPStructure for scans/images.

Strategy:
  - DOCX: python-docx direct extraction (text + tables)
  - PDF:  pymupdf direct extraction, scan pages fall back to PPStructure
  - Images (png/jpg/bmp/tiff): PPStructure
"""

from __future__ import annotations

import base64
import json
import logging
import re
from pathlib import Path

import httpx

from app.config import settings
from app.extraction.base import (
    BBox,
    OCRDetailedResult,
    OCRPageResult,
    OCRTextBlock,
)
from app.extraction.ocr.base import OCRProvider

logger = logging.getLogger(__name__)

# Minimum characters to consider a page as having extractable text.
_SCAN_PAGE_THRESHOLD = 20


class PaddleOCRProvider(OCRProvider):
    """Real OCR provider using direct extraction + PPStructure."""

    def extract_detailed(self, file_path: str, file_type: str) -> OCRDetailedResult:
        ext = Path(file_path).suffix.lower().lstrip(".")
        if ext == "jpeg":
            ext = "jpg"
        if ext == "tif":
            ext = "tiff"

        if file_type == "docx" or ext == "docx":
            return self._extract_docx(file_path)
        if file_type == "pdf" or ext == "pdf":
            return self._extract_pdf(file_path)
        # Images go straight to PPStructure
        return self._extract_image(file_path)

    # ------------------------------------------------------------------
    # DOCX direct extraction
    # ------------------------------------------------------------------

    def _extract_docx(self, file_path: str) -> OCRDetailedResult:
        from docx import Document

        doc = Document(file_path)
        blocks: list[OCRTextBlock] = []
        sort = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            sort += 1
            is_heading = para.style is not None and "heading" in (
                para.style.name or ""
            ).lower()
            blocks.append(OCRTextBlock(
                block_type="title" if is_heading else "text",
                text=text,
                bbox=None,
                confidence=1.0,
                sort_order=sort,
            ))

        for table in doc.tables:
            md = self._table_to_markdown(table)
            if md:
                sort += 1
                blocks.append(OCRTextBlock(
                    block_type="table",
                    text=md,
                    bbox=None,
                    confidence=1.0,
                    sort_order=sort,
                ))

        page = OCRPageResult(
            page_no=1,
            width=None,
            height=None,
            confidence=1.0,
            blocks=blocks,
        )
        return OCRDetailedResult(pages=[page], provider="paddle_docx")

    @staticmethod
    def _table_to_markdown(table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if not rows:
            return ""
        col_count = len(table.rows[0].cells)
        sep = "| " + " | ".join(["---"] * col_count) + " |"
        return rows[0] + "\n" + sep + "\n" + "\n".join(rows[1:])

    # ------------------------------------------------------------------
    # PDF direct extraction (+ scan page fallback)
    # ------------------------------------------------------------------

    def _extract_pdf(self, file_path: str) -> OCRDetailedResult:
        import fitz

        doc = fitz.open(file_path)
        pages: list[OCRPageResult] = []

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_no = page_idx + 1
            blocks: list[OCRTextBlock] = []
            sort = 0

            # Extract tables first
            table_regions: list[tuple[float, float, float, float]] = []
            try:
                for tab in page.find_tables():
                    md = tab.to_markdown()
                    if md and md.strip():
                        sort += 1
                        bbox = tab.bbox
                        table_regions.append(bbox)
                        blocks.append(OCRTextBlock(
                            block_type="table",
                            text=md.strip(),
                            bbox=BBox(x1=bbox[0], y1=bbox[1], x2=bbox[2], y2=bbox[3]),
                            confidence=1.0,
                            sort_order=sort,
                        ))
            except Exception:
                logger.debug("Table extraction failed on page %d", page_no, exc_info=True)

            # Extract text blocks
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            raw_text_len = 0
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                parts: list[str] = []
                max_font_size: float = 0.0
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        parts.append(span.get("text", ""))
                        # Capture largest font size in this block
                        sz = span.get("size", 0)
                        if sz > max_font_size:
                            max_font_size = float(sz)
                    parts.append("\n")
                text = "".join(parts).strip()
                if not text:
                    continue
                raw_text_len += len(text)

                bb = block.get("bbox")
                if bb and self._bbox_overlaps_table(bb, table_regions):
                    continue

                sort += 1
                bbox_obj = None
                if bb:
                    bbox_obj = BBox(x1=bb[0], y1=bb[1], x2=bb[2], y2=bb[3])
                block_type = self._guess_block_type(block)
                blocks.append(OCRTextBlock(
                    block_type=block_type,
                    text=text,
                    bbox=bbox_obj,
                    confidence=1.0,
                    sort_order=sort,
                    font_size=max_font_size if max_font_size > 0 else None,
                ))

            # Scan page detection
            if raw_text_len < _SCAN_PAGE_THRESHOLD:
                logger.info(
                    "Page %d appears scanned (%d chars), invoking PPStructure",
                    page_no, raw_text_len,
                )
                pixmap = page.get_pixmap(dpi=200)
                img_bytes = pixmap.tobytes("png")
                ocr_blocks = self._call_ppstructure(img_bytes, page_no, sort_offset=sort)
                if ocr_blocks:
                    blocks.extend(ocr_blocks)

            # Sort by position
            blocks.sort(key=lambda b: (
                (b.bbox.y1 if b.bbox else 0),
                (b.bbox.x1 if b.bbox else 0),
            ))
            for i, b in enumerate(blocks):
                b.sort_order = i + 1

            pages.append(OCRPageResult(
                page_no=page_no,
                width=int(page.rect.width),
                height=int(page.rect.height),
                confidence=1.0,
                blocks=blocks,
            ))

        doc.close()
        return OCRDetailedResult(pages=pages, provider="paddle_pdf")

    @staticmethod
    def _guess_block_type(block: dict) -> str:
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if span.get("size", 0) >= 14:
                    return "title"
        return "text"

    @staticmethod
    def _bbox_overlaps_table(bbox: tuple, table_regions: list) -> bool:
        x0, y0, x1, y1 = bbox
        for tx0, ty0, tx1, ty1 in table_regions:
            if x0 < tx1 and x1 > tx0 and y0 < ty1 and y1 > ty0:
                return True
        return False

    # ------------------------------------------------------------------
    # Image / PPStructure path
    # ------------------------------------------------------------------

    def _extract_image(self, file_path: str) -> OCRDetailedResult:
        img_bytes = Path(file_path).read_bytes()
        blocks = self._call_ppstructure(img_bytes, page_no=1)
        return OCRDetailedResult(
            pages=[OCRPageResult(
                page_no=1,
                width=None,
                height=None,
                confidence=blocks[0].confidence if blocks else 0.0,
                blocks=blocks,
            )],
            provider="paddle_ppstructure",
        )

    def _call_ppstructure(
        self,
        img_bytes: bytes,
        page_no: int,
        sort_offset: int = 0,
    ) -> list[OCRTextBlock]:
        """Call PPStructure HTTP service and parse results."""
        b64 = base64.b64encode(img_bytes).decode("ascii")
        payload = {"file": b64, "fileType": 1}

        url = settings.ppstructure_url
        timeout = settings.ocr_timeout

        response = self._http_post(url, payload, timeout)
        if response is None:
            raise RuntimeError(f"PPStructure call failed for page {page_no}")

        try:
            data = response.json()
        except (json.JSONDecodeError, ValueError) as exc:
            raise RuntimeError(f"PPStructure returned invalid JSON: {exc}") from exc

        regions = self._normalize_ppstructure_response(data)

        blocks: list[OCRTextBlock] = []
        for idx, region in enumerate(regions):
            rtype = region.get("type") or region.get("block_label") or "text"
            text = region.get("text") or region.get("block_content") or ""
            confidence = region.get("confidence", 0.8)
            raw_bbox = region.get("bbox") or region.get("block_bbox")
            raw_sort_order = region.get("block_order")

            if rtype == "table":
                text = self._table_html_to_markdown(text)

            if not text:
                continue

            bbox_obj = None
            if raw_bbox and len(raw_bbox) >= 4:
                bbox_obj = BBox(
                    x1=raw_bbox[0], y1=raw_bbox[1],
                    x2=raw_bbox[2], y2=raw_bbox[3],
                )

            blocks.append(OCRTextBlock(
                block_type=rtype,
                text=text.strip(),
                bbox=bbox_obj,
                confidence=confidence,
                sort_order=sort_offset + (raw_sort_order if isinstance(raw_sort_order, int) else idx + 1),
            ))

        if not blocks:
            raise RuntimeError(f"PPStructure returned no text blocks for page {page_no}")

        return blocks

    @staticmethod
    def _normalize_ppstructure_response(data) -> list[dict]:
        if isinstance(data, list):
            return data
        if isinstance(data, dict):
            error_code = data.get("errorCode")
            if error_code not in (None, 0):
                log_id = data.get("logId", "-")
                error_msg = data.get("errorMsg", "unknown PPStructure error")
                raise RuntimeError(f"PPStructure error {error_code} ({log_id}): {error_msg}")

            result = data.get("result")
            if isinstance(result, dict):
                layout_results = result.get("layoutParsingResults")
                if isinstance(layout_results, list):
                    regions: list[dict] = []
                    markdown_texts: list[str] = []
                    for page_result in layout_results:
                        if not isinstance(page_result, dict):
                            continue
                        pruned = page_result.get("prunedResult")
                        if isinstance(pruned, dict):
                            parsing_list = pruned.get("parsing_res_list")
                            if isinstance(parsing_list, list):
                                regions.extend(item for item in parsing_list if isinstance(item, dict))
                        markdown = page_result.get("markdown")
                        if isinstance(markdown, dict):
                            text = markdown.get("text")
                            if isinstance(text, str) and text.strip():
                                markdown_texts.append(text.strip())
                    if regions:
                        return regions
                    if markdown_texts:
                        return [{
                            "block_label": "text",
                            "block_content": "\n".join(markdown_texts),
                            "block_order": 1,
                        }]

            # Handle layout-aware response (PPStructureV3+)
            layout = data.get("layout")
            if layout and isinstance(layout, list):
                logger.debug("PPStructure returned layout-paragraphs response (%d paragraphs)", len(layout))
                return layout
            for key in ("result", "results", "regions", "layout_parsing_result"):
                if key in data:
                    val = data[key]
                    if isinstance(val, list):
                        return val
            if all(k.isdigit() for k in data.keys()):
                return list(data.values())
            logger.debug("PPStructure unexpected response keys: %s", list(data.keys()))
        return []

    @staticmethod
    def _table_html_to_markdown(text: str) -> str:
        if "<" not in text:
            return text
        rows: list[str] = []
        for tr_match in re.finditer(r"<tr[^>]*>(.*?)</tr>", text, re.DOTALL | re.IGNORECASE):
            cells = re.findall(
                r"<t[hd][^>]*>(.*?)</t[hd]>",
                tr_match.group(1),
                re.DOTALL | re.IGNORECASE,
            )
            if cells:
                cleaned = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
                rows.append("| " + " | ".join(cleaned) + " |")
        if not rows:
            return text
        col_count = rows[0].count("|") - 1
        sep = "| " + " | ".join(["---"] * col_count) + " |"
        return rows[0] + "\n" + sep + "\n" + "\n".join(rows[1:])

    # ------------------------------------------------------------------
    # HTTP helper with retry
    # ------------------------------------------------------------------

    @staticmethod
    def _http_post(
        url: str, payload: dict, timeout: int, retries: int = 2,
    ) -> httpx.Response | None:
        for attempt in range(1 + retries):
            try:
                with httpx.Client(timeout=timeout) as client:
                    resp = client.post(url, json=payload)
                    if resp.status_code >= 400:
                        try:
                            error_payload = resp.json()
                        except (json.JSONDecodeError, ValueError):
                            error_payload = None
                        if isinstance(error_payload, dict) and "errorCode" in error_payload:
                            return resp
                    resp.raise_for_status()
                    return resp
            except httpx.HTTPError as exc:
                logger.warning(
                    "HTTP request failed (attempt %d/%d): %s",
                    attempt + 1, 1 + retries, exc,
                )
        return None
