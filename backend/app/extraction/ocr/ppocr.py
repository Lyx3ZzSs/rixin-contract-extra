"""PP-OCR provider -- direct text extraction + fast OCR for scans/images.

Strategy:
  - DOCX: python-docx direct extraction (text + tables)
  - PDF:  pymupdf direct extraction, scan pages fall back to PP-OCR
  - Images (png/jpg/bmp/tiff): PP-OCR
"""

from __future__ import annotations

import base64
import json
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any

import httpx

from app.config import settings
from app.extraction.base import BBox, OCRDetailedResult, OCRPageResult, OCRTextBlock
from app.extraction.ocr.base import OCRProvider

logger = logging.getLogger(__name__)

_SCAN_PAGE_THRESHOLD = 20


class PPOCRProvider(OCRProvider):
    """Fast OCR provider using direct extraction plus a PP-OCR HTTP service."""

    def extract_detailed(self, file_path: str, file_type: str) -> OCRDetailedResult:
        ext = Path(file_path).suffix.lower().lstrip(".")
        if ext == "jpeg":
            ext = "jpg"
        if ext == "tif":
            ext = "tiff"

        if file_type == "docx" or ext == "docx":
            return self._extract_docx(file_path)
        if file_type == "pdf" or ext == "pdf":
            return self._extract_pdf(file_path)
        return self._extract_image(file_path)

    def _extract_docx(self, file_path: str) -> OCRDetailedResult:
        from docx import Document

        doc = Document(file_path)
        blocks: list[OCRTextBlock] = []
        sort = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            sort += 1
            is_heading = para.style is not None and "heading" in (para.style.name or "").lower()
            blocks.append(OCRTextBlock(
                block_type="title" if is_heading else "text",
                text=text,
                bbox=None,
                confidence=1.0,
                sort_order=sort,
            ))

        for table in doc.tables:
            md = self._table_to_markdown(table)
            if md:
                sort += 1
                blocks.append(OCRTextBlock(
                    block_type="table",
                    text=md,
                    bbox=None,
                    confidence=1.0,
                    sort_order=sort,
                ))

        return OCRDetailedResult(
            pages=[OCRPageResult(page_no=1, blocks=blocks, confidence=1.0)],
            provider="ppocr_docx",
        )

    @staticmethod
    def _table_to_markdown(table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if not rows:
            return ""
        col_count = len(table.rows[0].cells)
        sep = "| " + " | ".join(["---"] * col_count) + " |"
        return rows[0] + "\n" + sep + "\n" + "\n".join(rows[1:])

    def _extract_pdf(self, file_path: str) -> OCRDetailedResult:
        import fitz

        doc = fitz.open(file_path)
        try:
            text_pages = self._extract_pdf_text_pages(doc)
            text_len = len("\n".join(page.full_text for page in text_pages).strip())
            text_page_count = sum(1 for page in text_pages if page.full_text.strip())
            page_count = len(text_pages)
            text_page_ratio = text_page_count / page_count if page_count else 0

            if (
                text_len >= settings.ppocr_pdf_text_min_chars
                and text_page_ratio >= settings.ppocr_pdf_text_page_ratio
            ):
                logger.info(
                    "PDF classified as text: pages=%d text_pages=%d text_len=%d ratio=%.2f",
                    page_count, text_page_count, text_len, text_page_ratio,
                )
                return OCRDetailedResult(pages=text_pages, provider="ppocr_pdf_text")

            logger.info(
                "PDF classified as OCR-needed: pages=%d text_pages=%d text_len=%d ratio=%.2f",
                page_count, text_page_count, text_len, text_page_ratio,
            )

            try:
                whole_pages = self._call_ppocr_pdf(Path(file_path).read_bytes(), text_pages)
                if self._has_any_text(whole_pages):
                    return OCRDetailedResult(pages=whole_pages, provider="ppocr_pdf_whole")
                logger.warning("Whole-PDF PP-OCR returned no text; falling back to page-image OCR")
            except RuntimeError as exc:
                logger.warning("Whole-PDF PP-OCR failed, falling back to page-image OCR: %s", exc)

            fallback_pages = self._extract_pdf_page_images(doc, text_pages)
            return OCRDetailedResult(pages=fallback_pages, provider="ppocr_pdf_page_fallback")
        finally:
            doc.close()

    def _extract_pdf_text_pages(self, doc: Any) -> list[OCRPageResult]:
        import fitz

        pages: list[OCRPageResult] = []
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_no = page_idx + 1
            blocks: list[OCRTextBlock] = []
            sort = 0

            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                parts: list[str] = []
                max_font_size = 0.0
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        parts.append(span.get("text", ""))
                        size = float(span.get("size", 0) or 0)
                        max_font_size = max(max_font_size, size)
                    parts.append("\n")
                text = "".join(parts).strip()
                if not text:
                    continue

                sort += 1
                bbox_obj = None
                bbox = block.get("bbox")
                if bbox:
                    bbox_obj = BBox(x1=bbox[0], y1=bbox[1], x2=bbox[2], y2=bbox[3])
                blocks.append(OCRTextBlock(
                    block_type=self._guess_block_type(block),
                    text=text,
                    bbox=bbox_obj,
                    confidence=1.0,
                    sort_order=sort,
                    font_size=max_font_size if max_font_size > 0 else None,
                ))

            self._sort_blocks(blocks)
            pages.append(OCRPageResult(
                page_no=page_no,
                width=int(page.rect.width),
                height=int(page.rect.height),
                confidence=1.0,
                blocks=blocks,
            ))
        return pages

    def _extract_pdf_page_images(
        self,
        doc: Any,
        text_pages: list[OCRPageResult],
    ) -> list[OCRPageResult]:
        page_jobs: list[tuple[int, bytes, int]] = []
        pages = [
            OCRPageResult(
                page_no=page.page_no,
                width=page.width,
                height=page.height,
                confidence=page.confidence,
                blocks=list(page.blocks),
            )
            for page in text_pages
        ]

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_result = pages[page_idx]
            raw_text_len = len(page_result.full_text.strip())
            if raw_text_len >= _SCAN_PAGE_THRESHOLD:
                continue
            logger.info(
                "Page %d appears scanned (%d chars), invoking page-image PP-OCR",
                page_result.page_no, raw_text_len,
            )
            pixmap = page.get_pixmap(dpi=settings.ppocr_pdf_dpi)
            page_jobs.append((page_idx, pixmap.tobytes("png"), len(page_result.blocks)))

        if not page_jobs:
            return pages

        errors: list[str] = []
        max_workers = max(1, min(settings.ppocr_page_concurrency, len(page_jobs)))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_map = {
                executor.submit(
                    self._call_ppocr,
                    img_bytes,
                    pages[page_idx].page_no,
                    sort_offset,
                    True,
                ): page_idx
                for page_idx, img_bytes, sort_offset in page_jobs
            }
            for future in as_completed(future_map):
                page_idx = future_map[future]
                page_result = pages[page_idx]
                try:
                    page_result.blocks.extend(future.result())
                    self._sort_blocks(page_result.blocks)
                except RuntimeError as exc:
                    errors.append(f"page {page_result.page_no}: {exc}")
                    logger.warning("Page-image PP-OCR failed for page %d: %s", page_result.page_no, exc)

        if not self._has_any_text(pages) and errors:
            raise RuntimeError("; ".join(errors[:3]))
        return pages

    @staticmethod
    def _guess_block_type(block: dict) -> str:
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if span.get("size", 0) >= 14:
                    return "title"
        return "text"

    def _extract_image(self, file_path: str) -> OCRDetailedResult:
        blocks = self._call_ppocr(Path(file_path).read_bytes(), page_no=1)
        return OCRDetailedResult(
            pages=[OCRPageResult(
                page_no=1,
                width=None,
                height=None,
                confidence=blocks[0].confidence if blocks else 0.0,
                blocks=blocks,
            )],
            provider="ppocr",
        )

    def _call_ppocr(
        self,
        img_bytes: bytes,
        page_no: int,
        sort_offset: int = 0,
        allow_empty: bool = False,
    ) -> list[OCRTextBlock]:
        payload = self._build_ppocr_payload(img_bytes, file_type=1)
        response = self._http_post(settings.ppocr_url, payload, settings.ppocr_timeout)
        if response is None:
            raise RuntimeError(f"PP-OCR call failed for page {page_no}")

        try:
            data = response.json()
        except (json.JSONDecodeError, ValueError) as exc:
            raise RuntimeError(f"PP-OCR returned invalid JSON: {exc}") from exc

        blocks = self._normalize_ppocr_response(data, sort_offset=sort_offset)
        if not blocks:
            logger.warning(
                "PP-OCR returned no text blocks for page %d; response_summary=%s",
                page_no, self._response_summary(data),
            )
            if allow_empty:
                return []
            raise RuntimeError(f"PP-OCR returned no text blocks for page {page_no}")
        return blocks

    def _call_ppocr_pdf(
        self,
        pdf_bytes: bytes,
        fallback_pages: list[OCRPageResult],
    ) -> list[OCRPageResult]:
        payload = self._build_ppocr_payload(pdf_bytes, file_type=0)
        response = self._http_post(settings.ppocr_url, payload, settings.ppocr_timeout)
        if response is None:
            raise RuntimeError("PP-OCR whole-PDF call failed")

        try:
            data = response.json()
        except (json.JSONDecodeError, ValueError) as exc:
            raise RuntimeError(f"PP-OCR returned invalid JSON: {exc}") from exc

        pages = self._normalize_ppocr_pdf_response(data, fallback_pages)
        if not self._has_any_text(pages):
            logger.warning(
                "Whole-PDF PP-OCR returned no text blocks; response_summary=%s",
                self._response_summary(data),
            )
        return pages

    @classmethod
    def _normalize_ppocr_response(cls, data: Any, sort_offset: int = 0) -> list[OCRTextBlock]:
        if isinstance(data, dict):
            error_code = data.get("errorCode")
            if error_code not in (None, 0):
                log_id = data.get("logId", "-")
                error_msg = data.get("errorMsg", "unknown PP-OCR error")
                raise RuntimeError(f"PP-OCR error {error_code} ({log_id}): {error_msg}")
            data = cls._extract_result_payload(data)

        rows = cls._flatten_rows(data)
        blocks: list[OCRTextBlock] = []
        for index, row in enumerate(rows, start=1):
            parsed = cls._parse_ocr_row(row)
            if parsed is None:
                continue
            text, bbox, confidence = parsed
            text = text.strip()
            if not text:
                continue
            blocks.append(OCRTextBlock(
                block_type="text",
                text=text,
                bbox=bbox,
                confidence=confidence,
                sort_order=sort_offset + index,
            ))

        blocks.sort(key=lambda b: ((b.bbox.y1 if b.bbox else 0), (b.bbox.x1 if b.bbox else 0)))
        for index, block in enumerate(blocks, start=1):
            block.sort_order = sort_offset + index
        return blocks

    @classmethod
    def _normalize_ppocr_pdf_response(
        cls,
        data: Any,
        fallback_pages: list[OCRPageResult],
    ) -> list[OCRPageResult]:
        if isinstance(data, dict):
            error_code = data.get("errorCode")
            if error_code not in (None, 0):
                log_id = data.get("logId", "-")
                error_msg = data.get("errorMsg", "unknown PP-OCR error")
                raise RuntimeError(f"PP-OCR error {error_code} ({log_id}): {error_msg}")

        result = data.get("result") if isinstance(data, dict) else None
        if not isinstance(result, dict) or not isinstance(result.get("ocrResults"), list):
            blocks = cls._normalize_ppocr_response(data)
            page = fallback_pages[0] if fallback_pages else OCRPageResult(page_no=1)
            return [
                OCRPageResult(
                    page_no=page.page_no,
                    width=page.width,
                    height=page.height,
                    confidence=1.0,
                    blocks=blocks,
                ),
            ]

        dimensions = cls._extract_pdf_dimensions(result.get("dataInfo"), fallback_pages)
        pages: list[OCRPageResult] = []
        for index, ocr_result in enumerate(result.get("ocrResults", [])):
            fallback = fallback_pages[index] if index < len(fallback_pages) else None
            page_no = fallback.page_no if fallback else index + 1
            width, height = dimensions[index] if index < len(dimensions) else (
                fallback.width if fallback else None,
                fallback.height if fallback else None,
            )
            payload = ocr_result
            if isinstance(ocr_result, dict):
                payload = ocr_result.get("prunedResult", ocr_result)
            blocks = cls._normalize_ppocr_response(payload)
            pages.append(OCRPageResult(
                page_no=page_no,
                width=width,
                height=height,
                confidence=1.0,
                blocks=blocks,
            ))

        return pages

    @staticmethod
    def _extract_result_payload(data: dict) -> Any:
        for key in ("result", "results", "data", "ocr_results", "ocrResults"):
            value = data.get(key)
            if value is not None:
                return PPOCRProvider._extract_structured_payload(value)
        return PPOCRProvider._extract_structured_payload(data)

    @staticmethod
    def _extract_structured_payload(data: Any) -> Any:
        if isinstance(data, dict):
            if isinstance(data.get("ocrResults"), list):
                rows: list[Any] = []
                for item in data.get("ocrResults", []):
                    payload = item
                    if isinstance(item, dict):
                        payload = item.get("prunedResult", item)
                    extracted = PPOCRProvider._extract_structured_payload(payload)
                    rows.extend(PPOCRProvider._flatten_rows(extracted))
                return rows

            if "prunedResult" in data:
                return PPOCRProvider._extract_structured_payload(data.get("prunedResult"))

            rec_texts = data.get("rec_texts") or data.get("texts")
            if isinstance(rec_texts, list):
                boxes = (
                    data.get("rec_polys")
                    or data.get("dt_polys")
                    or data.get("rec_boxes")
                    or data.get("boxes")
                    or []
                )
                scores = data.get("rec_scores") or data.get("scores") or []
                return [
                    {
                        "text": text,
                        "bbox": boxes[idx] if idx < len(boxes) else None,
                        "confidence": scores[idx] if idx < len(scores) else 0.8,
                    }
                    for idx, text in enumerate(rec_texts)
                ]
        return data

    @classmethod
    def _flatten_rows(cls, data: Any) -> list[Any]:
        if isinstance(data, list):
            if len(data) == 1 and isinstance(data[0], list):
                inner = data[0]
                if (
                    len(inner) >= 2
                    and isinstance(inner[1], (list, tuple))
                    and len(inner[1]) >= 1
                    and isinstance(inner[1][0], str)
                ):
                    return data
                return cls._flatten_rows(data[0])
            return data
        if isinstance(data, dict):
            return [data]
        return []

    @classmethod
    def _parse_ocr_row(cls, row: Any) -> tuple[str, BBox | None, float] | None:
        if isinstance(row, dict):
            text = row.get("text") or row.get("transcription") or row.get("rec_text") or row.get("value")
            bbox = cls._bbox_from_raw(row.get("bbox") or row.get("box") or row.get("points") or row.get("polygon"))
            confidence = row.get("confidence", row.get("score", row.get("rec_score", 0.8)))
            return (str(text or ""), bbox, cls._safe_float(confidence, 0.8))

        if isinstance(row, (list, tuple)):
            # PaddleOCR common format: [box, (text, score)]
            if len(row) >= 2 and isinstance(row[1], (list, tuple)) and len(row[1]) >= 1:
                bbox = cls._bbox_from_raw(row[0])
                text = str(row[1][0] or "")
                confidence = cls._safe_float(row[1][1] if len(row[1]) > 1 else 0.8, 0.8)
                return (text, bbox, confidence)
            # Simpler format: [text, score] or [text]
            if row and isinstance(row[0], str):
                confidence = cls._safe_float(row[1] if len(row) > 1 else 0.8, 0.8)
                return (row[0], None, confidence)
        return None

    @staticmethod
    def _bbox_from_raw(raw: Any) -> BBox | None:
        if not raw:
            return None
        if isinstance(raw, dict):
            values = [raw.get(k) for k in ("x1", "y1", "x2", "y2")]
            if all(v is not None for v in values):
                return BBox(x1=float(values[0]), y1=float(values[1]), x2=float(values[2]), y2=float(values[3]))
            return None
        if isinstance(raw, (list, tuple)):
            if len(raw) >= 4 and all(isinstance(v, (int, float)) for v in raw[:4]):
                return BBox(x1=float(raw[0]), y1=float(raw[1]), x2=float(raw[2]), y2=float(raw[3]))
            if raw and all(isinstance(point, (list, tuple)) and len(point) >= 2 for point in raw):
                xs = [float(point[0]) for point in raw]
                ys = [float(point[1]) for point in raw]
                return BBox(x1=min(xs), y1=min(ys), x2=max(xs), y2=max(ys))
        return None

    @staticmethod
    def _safe_float(value: Any, default: float) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return default

    @staticmethod
    def _build_ppocr_payload(file_bytes: bytes, file_type: int) -> dict:
        return {
            "file": base64.b64encode(file_bytes).decode("ascii"),
            "fileType": file_type,
            "visualize": False,
            "returnWordBox": False,
        }

    @staticmethod
    def _extract_pdf_dimensions(
        data_info: Any,
        fallback_pages: list[OCRPageResult],
    ) -> list[tuple[int | None, int | None]]:
        if isinstance(data_info, dict) and isinstance(data_info.get("pages"), list):
            return [
                (
                    int(page.get("width")) if page.get("width") is not None else None,
                    int(page.get("height")) if page.get("height") is not None else None,
                )
                for page in data_info.get("pages", [])
                if isinstance(page, dict)
            ]
        if isinstance(data_info, dict) and data_info.get("width") is not None and data_info.get("height") is not None:
            return [(int(data_info.get("width")), int(data_info.get("height")))]
        return [(page.width, page.height) for page in fallback_pages]

    @staticmethod
    def _sort_blocks(blocks: list[OCRTextBlock]) -> None:
        blocks.sort(key=lambda b: ((b.bbox.y1 if b.bbox else 0), (b.bbox.x1 if b.bbox else 0)))
        for index, block in enumerate(blocks, start=1):
            block.sort_order = index

    @staticmethod
    def _has_any_text(pages: list[OCRPageResult]) -> bool:
        return any(page.full_text.strip() for page in pages)

    @staticmethod
    def _response_summary(data: Any) -> dict:
        if not isinstance(data, dict):
            return {"type": type(data).__name__}

        result = data.get("result")
        summary: dict[str, Any] = {
            "top_keys": sorted(data.keys()),
            "result_type": type(result).__name__,
        }
        if isinstance(result, dict):
            ocr_results = result.get("ocrResults")
            summary["result_keys"] = sorted(result.keys())
            summary["ocrResults_count"] = len(ocr_results) if isinstance(ocr_results, list) else None
            if isinstance(ocr_results, list) and ocr_results:
                first = ocr_results[0]
                if isinstance(first, dict):
                    pruned = first.get("prunedResult")
                    summary["first_ocr_result_keys"] = sorted(first.keys())
                    if isinstance(pruned, dict):
                        summary["first_prunedResult_keys"] = sorted(pruned.keys())
        return summary

    @staticmethod
    def _http_post(url: str, payload: dict, timeout: int, retries: int = 2) -> httpx.Response | None:
        if not url:
            raise RuntimeError("PPOCR_URL is not configured")
        for attempt in range(1 + retries):
            try:
                with httpx.Client(timeout=timeout) as client:
                    resp = client.post(url, json=payload)
                    if resp.status_code >= 400:
                        try:
                            error_payload = resp.json()
                        except (json.JSONDecodeError, ValueError):
                            error_payload = None
                        if isinstance(error_payload, dict) and "errorCode" in error_payload:
                            return resp
                    resp.raise_for_status()
                    return resp
            except httpx.HTTPError as exc:
                logger.warning(
                    "PP-OCR request failed (attempt %d/%d): %s",
                    attempt + 1, 1 + retries, exc,
                )
        return None
